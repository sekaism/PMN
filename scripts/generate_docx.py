#!/usr/bin/env python3
"""
Generate formatted DOCX manuscript with embedded figures.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

BASE = "/workspaces/platform/MN"
FIGURES = os.path.join(BASE, "figures")

doc = Document()

# ================================================================
# Styles
# ================================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    heading_style = doc.styles[f'Heading {level}']
    heading_style.font.color.rgb = RGBColor(0, 51, 102)
    heading_style.font.name = 'Times New Roman'
    heading_style.font.bold = True

# ================================================================
# Title Page
# ================================================================
for _ in range(4):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Integrative Network Medicine Framework Reveals Multi-Target Mechanisms of Rituximab in Primary Membranous Nephropathy')
run.font.size = Pt(18)
run.font.bold = True
run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('[Authors to be added]')
run.font.size = Pt(12)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run('[Affiliations to be added]')
run.font.size = Pt(10)
run.font.italic = True

doc.add_paragraph()

corr = doc.add_paragraph()
corr.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = corr.add_run('Corresponding author: [To be added]')
run.font.size = Pt(10)

doc.add_page_break()

# ================================================================
# Abstract
# ================================================================
doc.add_heading('Abstract', level=1)

abstract_text = (
    "Background: Rituximab (RTX) has emerged as a first-line therapy for primary membranous "
    "nephropathy (PMN), yet approximately 35-40% of patients do not achieve clinical remission. "
    "The mechanisms underlying RTX efficacy extend beyond B-cell depletion, but a systematic "
    "characterization of RTX's multi-target effects on the PMN molecular network remains lacking.\n\n"
    "Methods: We developed a multi-layer network medicine framework integrating: (1) transcriptomic "
    "differential expression analysis of 87 PMN versus 11 normal glomerular samples; (2) GWAS risk "
    "loci (PLA2R1, NFKB1, IRF4, HLA); (3) drug-target interaction networks for RTX, tacrolimus (TAC), "
    "and cyclophosphamide (CTX); and (4) a PPI network comprising 2,176 nodes and 3,098 edges.\n\n"
    "Results: We identified 3,063 differentially expressed genes. RTX targets are significantly closer "
    "to the PMN disease module (z = -5.49, p < 0.001) compared with TAC (z = -1.89) and CTX (z = -2.54). "
    "RTX and TAC show complementary exposure pattern (separation s = 1.12), and RTX targets are enriched "
    "in B cell (89%) and NK cell (33%) markers, while TAC targets T cells (11%).\n\n"
    "Conclusion: This study provides the first systematic network medicine characterization of RTX's "
    "therapeutic mechanisms in PMN, demonstrating multi-target effects beyond CD20-mediated B-cell depletion."
)
p = doc.add_paragraph(abstract_text)
p.paragraph_format.first_line_indent = Cm(0)

doc.add_paragraph()
kw = doc.add_paragraph()
run = kw.add_run('Keywords: ')
run.font.bold = True
run = kw.add_run('membranous nephropathy; rituximab; network medicine; network pharmacology; drug combination; transcriptomics')

doc.add_page_break()

# ================================================================
# 1. Introduction
# ================================================================
doc.add_heading('1. Introduction', level=1)

sections = [
    ("", [
        "Primary membranous nephropathy (PMN) is an autoimmune kidney disease and the most common "
        "cause of nephrotic syndrome in adults, accounting for approximately 25-40% of cases [1,2]. "
        "The disease is characterized by the deposition of immune complexes in the subepithelial space "
        "of the glomerular basement membrane, leading to podocyte injury and proteinuria. The identification "
        "of autoantibodies against the M-type phospholipase A2 receptor (PLA2R) as the primary pathogenic "
        "antigen in approximately 70-80% of patients revolutionized our understanding of PMN as a "
        "B-cell-mediated autoimmune disorder [3,4].",

        "Rituximab (RTX), a chimeric monoclonal antibody targeting CD20 on B lymphocytes, has become "
        "a first-line therapy for PMN based on landmark randomized controlled trials including MENTOR "
        "and RI-CYCLO [5,6]. The 2021 KDIGO guidelines recommend RTX as a first-line treatment for "
        "PMN patients at intermediate-to-high risk of progression [7]. However, approximately 35-40% "
        "of patients do not achieve clinical remission following RTX therapy, and the mechanisms "
        "underlying treatment resistance remain poorly understood [8,9].",

        "While RTX's canonical mechanism involves B-cell depletion leading to reduced anti-PLA2R "
        "antibody production, emerging evidence suggests additional therapeutic actions beyond B-cell "
        "targeting. Recent studies have demonstrated that RTX significantly modulates T lymphocyte "
        "subsets, including restoration of regulatory T cell (Treg) levels and normalization of the "
        "CD4+/CD8+ ratio, independently of B-cell depletion [10]. This T-cell modulatory effect "
        "appears to be mediated through NK cell-dependent TGF{\\beta} signaling during antibody-dependent "
        "cellular cytotoxicity (ADCC) [11]. Furthermore, RTX engages multiple effector mechanisms "
        "including complement-dependent cytotoxicity (CDC) through C1q binding, and Fc{\\gamma} "
        "receptor-mediated effects on innate immune cells [12].",

        "Traditional network pharmacology approaches have been applied to study herbal medicines "
        "in PMN [13,14], but these methods typically rely on simple target overlap analysis rather "
        "than quantitative network-based measurements. More fundamentally, no systematic network "
        "analysis has been applied to characterize RTX's multi-target therapeutic effects in PMN. "
        "The recently developed network medicine framework, pioneered by Barab{\\'a}si and "
        "colleagues, offers a powerful approach to quantify the topological relationship between "
        "drug targets and disease modules within the human protein-protein interactome [15,16]. "
        "Key metrics including network proximity (z-score) and drug-target separation (s-score) "
        "have been validated for predicting drug efficacy and drug combination synergy across "
        "multiple therapeutic areas [17,18].",

        "Here, we present the first multi-layer network medicine analysis of RTX in PMN. By "
        "integrating transcriptomic data from 87 PMN glomerular samples, GWAS risk loci, "
        "drug-target networks, and the human interactome, we quantify RTX's network proximity "
        "to the PMN disease module and benchmark it against tacrolimus (TAC) and cyclophosphamide "
        "(CTX). We further evaluate the network-based rationale for RTX+TAC combination therapy, "
        "which recent clinical network meta-analyses have identified as the most effective regimen "
        "for PMN [19]."
    ])
]

for _, paragraphs in sections:
    for text in paragraphs:
        p = doc.add_paragraph(text)
        p.paragraph_format.first_line_indent = Cm(0.75)

# ================================================================
# Figure 1: Workflow
# ================================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Figure1_Workflow.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.5))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figure 1. ')
run.font.bold = True
run = p.add_run('Multi-layer network medicine framework for analysis of rituximab in membranous nephropathy.')
run.font.italic = True

doc.add_page_break()

# ================================================================
# 2. Materials and Methods
# ================================================================
doc.add_heading('2. Materials and Methods', level=1)

doc.add_heading('2.1 Transcriptomic Data Acquisition and Processing', level=2)
p = doc.add_paragraph(
    "Genome-wide glomerular transcriptomic data were obtained from the Nephrotic Syndrome "
    "Study Network (NEPTUNE) cohort, accessed through the Gene Expression Omnibus under "
    "accession GSE108113 [20]. This dataset includes 280 microarray samples profiled on the "
    "Affymetrix Human Gene 2.1 ST Array platform, comprising 87 PMN, 11 normal living donor "
    "controls, and samples from other glomerular diseases. Only glomerular compartment samples "
    "were used. Differentially expressed genes (DEGs) between PMN and normal controls were "
    "identified using the limma package in R with empirical Bayes moderation [21]. Genes with "
    "Benjamini-Hochberg adjusted p-value < 0.05 were considered statistically significant."
)

doc.add_heading('2.2 PMN Disease Module Construction', level=2)
p = doc.add_paragraph(
    "The PMN disease module was constructed by integrating transcriptomic DEGs, established "
    "PMN GWAS risk genes including PLA2R1 (rs17831251, OR=2.25), NFKB1 (rs230540, OR=1.25), "
    "IRF4 (rs9405192, OR=1.29), and HLA-DQA1/DRB1 (rs9271573, OR=2.41) [22], and PPI "
    "network expansion using STRING database version 11.5 (confidence score >= 0.4) [23]. "
    "The largest connected component was retained as the final disease module."
)

doc.add_heading('2.3 Drug Target Collection', level=2)
p = doc.add_paragraph(
    "Drug targets were curated from DrugBank [24], STITCH [25], the Harmonizome collection [26], "
    "and published literature. For RTX, targets included MS4A1 (CD20), Fc{\\gamma} receptors "
    "(FCGR1A, FCGR2A, FCGR2B, FCGR3A), complement components (C1QA, C1QB, C1QC), B-cell "
    "receptor signaling pathway components (SYK, LYN, BTK, PLCG2, AKT1, NFKB1), and T-cell "
    "regulatory mediators (TGFB1, FOXP3, IL2RA). For TAC, targets included FKBP1A, calcineurin "
    "subunits (PPP3CA, PPP3CB, PPP3CC), and NFAT transcription factors (NFATC1-4). For CTX, "
    "targets included activation enzymes (CYP2B6, ALDH1A1) and DNA damage response genes "
    "(TP53, ATM, CASP3). All targets were expanded to include first-order interaction partners "
    "within the STRING PPI network."
)

doc.add_heading('2.4 Network Proximity and Separation Analysis', level=2)
p = doc.add_paragraph(
    "Network proximity between drug target sets (T) and disease module genes (D) was "
    "quantified using the average minimum shortest path length (AMSPL). Statistical "
    "significance was assessed through permutation tests (n=1,000 iterations). The "
    "network-based separation between two drug target sets A and B was calculated as: "
    "s_AB = <d_AB> - (<d_AA> + <d_BB>)/2, where <d_AB> is the mean shortest distance "
    "between targets of drug A and B [17,27]. The Complementary Exposure pattern was "
    "defined as both drug target sets being significantly proximal to the disease module "
    "(z < -1.5) while being separated from each other (s > 0.5) [27]."
)

doc.add_heading('2.5 Cell-Type Mapping', level=2)
p = doc.add_paragraph(
    "Single-cell RNA-seq data from 6 PMN and 2 normal kidney samples (GSE171458) were "
    "analyzed. Drug target expression was mapped to 12 kidney cell types using established "
    "cell-type marker genes from the KPMP kidney cell atlas and published literature. "
    "Cell-type enrichment was quantified as the percentage of cell-type markers present "
    "in each drug's target set."
)

doc.add_heading('2.6 Functional Enrichment Analysis', level=2)
p = doc.add_paragraph(
    "Gene Ontology Biological Process and KEGG pathway enrichment analyses were performed "
    "using Enrichr [28] with Benjamini-Hochberg correction (adjusted p < 0.05)."
)

# ================================================================
# 3. Results
# ================================================================
doc.add_heading('3. Results', level=1)

doc.add_heading('3.1 Transcriptomic Landscape of PMN Glomeruli', level=2)
p = doc.add_paragraph(
    "Comparative transcriptomic analysis of 87 PMN versus 11 normal glomerular samples "
    "identified 3,063 significantly differentially expressed genes at FDR < 0.05. Of these, "
    "1,601 genes were upregulated and 1,462 were downregulated in PMN. The most significantly "
    "upregulated genes included TP53, KEAP1, RRBP1, and MYDGF, while the most downregulated "
    "genes included SRSF5, MT1E, and IFIT1B. GO enrichment revealed significant enrichment "
    "for immune system activation, complement cascade, and inflammatory response pathways "
    "(all FDR < 0.05)."
)

doc.add_heading('3.2 PMN Disease Module', level=2)
p = doc.add_paragraph(
    "Integration of 3,063 DEGs with five GWAS risk genes followed by STRING PPI expansion "
    "produced a disease module comprising 2,176 proteins and 3,098 physical interactions "
    "(Figure 5). LCC analysis confirmed the module formed a significantly connected subnetwork "
    "(p < 0.001)."
)

doc.add_heading('3.3 Drug-Disease Network Proximity', level=2)
p = doc.add_paragraph(
    "RTX targets demonstrated the strongest proximity to the PMN disease module "
    "(z = -5.49, p < 0.001), indicating that RTX's mechanism is highly aligned with "
    "the molecular pathology of PMN (Figure 2). CTX targets showed moderate proximity "
    "(z = -2.54, p = 0.006), while TAC targets showed weaker proximity (z = -1.89, "
    "p = 0.029). The relative ranking (RTX > CTX > TAC) aligns with the KDIGO "
    "treatment algorithm [7]."
)

# Figure 2
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Figure2_Drug_Disease_Proximity.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.0))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figure 2. ')
run.font.bold = True
run = p.add_run('Drug-disease network proximity analysis.')
run.font.italic = True
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('Bar chart showing proximity z-scores for RTX (z = -5.49), CTX (z = -2.54), and TAC (z = -1.89). More negative values indicate closer network proximity to the PMN disease module. Dashed line indicates significance threshold (z = -1.5).')
run2.font.size = Pt(9)
run2.font.italic = True

doc.add_heading('3.4 Drug-Target Separation and Combination Prediction', level=2)
p = doc.add_paragraph(
    "All three drug pairs exhibited positive separation scores (Figure 3). Critically, "
    "both RTX and TAC demonstrated significant proximity to the PMN disease module while "
    "maintaining positive target separation (s = 1.12), matching the Complementary Exposure "
    "pattern associated with synergistic drug combination effects (Figure 4) [27]. "
    "This is consistent with recent clinical NMA evidence showing that RTX+TAC achieves "
    "the highest overall response rate (SUCRA = 93.5%) among all evaluated regimens [19]."
)

# Figure 3
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Figure3_Drug_Drug_Separation.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.0))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figure 3. ')
run.font.bold = True
run = p.add_run('Drug-drug target network separation.')
run.font.italic = True

# Figure 4
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Figure4_Complementary_Exposure.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.5))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figure 4. ')
run.font.bold = True
run = p.add_run('Complementary Exposure pattern between RTX and TAC.')
run.font.italic = True

doc.add_heading('3.5 Cell-Type Specificity', level=2)
p = doc.add_paragraph(
    "Cell-type mapping using scRNA-seq data (GSE171458) revealed complementary cellular "
    "targeting (Supplementary Figure S1). RTX targets were enriched in B cells "
    "(89% of cell-type markers), NK cells (33%), and macrophages (40%), consistent with "
    "its multi-layer mechanism of B-cell depletion, ADCC, and immunomodulation. TAC targets "
    "showed enrichment in T cells (11%), consistent with calcineurin-NFAT pathway inhibition. "
    "CTX targets showed no significant cell-type enrichment, reflecting its non-selective "
    "DNA-damaging mechanism."
)

# ================================================================
# Figure 5: PPI Network
# ================================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Figure5_PPI_Network.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.5))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figure 5. ')
run.font.bold = True
run = p.add_run('PPI network of the PMN disease module with drug target overlay.')
run.font.italic = True

doc.add_heading('3.6 Multi-Layer Mechanism of RTX', level=2)
p = doc.add_paragraph(
    "Integration of our network medicine findings with recent mechanistic studies reveals "
    "a three-layer model of RTX action (Figure 6). Layer 1 involves direct B-cell depletion "
    "through CD20 binding. Layer 2 encompasses Fc{\\gamma}R-mediated ADCC by NK cells and "
    "CDC via complement. Layer 3 involves NK cell-mediated TGF{\\beta} signaling that induces "
    "Treg expansion independently of B-cell depletion [11]. Our network analysis demonstrates "
    "that these three layers collectively position RTX targets in close proximity to the PMN "
    "disease module (z = -5.49)."
)

# Figure 6
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Figure6_Summary.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(5.5))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Figure 6. ')
run.font.bold = True
run = p.add_run('Summary of findings and multi-layer mechanism of RTX in PMN.')
run.font.italic = True

# ================================================================
# 4. Discussion
# ================================================================
doc.add_heading('4. Discussion', level=1)

discussion_paragraphs = [
    "This study presents the first comprehensive network medicine characterization of "
    "rituximab's therapeutic mechanisms in primary membranous nephropathy. Our multi-layer "
    "framework integrates transcriptomic, genetic, and pharmacological data to provide a "
    "systems-level understanding of how RTX targets the PMN disease network.",

    "The concept of network proximity provides a powerful quantitative metric for predicting "
    "therapeutic efficacy [17,18]. Our finding that RTX exhibits the strongest proximity to "
    "the PMN disease module (z = -5.49) provides quantitative validation for its first-line "
    "clinical status. Notably, RTX's proximity to GWAS-defined risk genes (z = -9.26) was "
    "markedly stronger than to transcriptomic DEGs alone, suggesting that RTX's therapeutic "
    "effect may be particularly relevant to the core genetic susceptibility pathways in PMN.",

    "The identification of a Complementary Exposure pattern between RTX and TAC provides "
    "a network-based mechanistic rationale for their synergistic combination. This pattern, "
    "previously validated across multiple therapeutic areas [17,27], indicates that the two "
    "drugs target distinct but disease-relevant network neighborhoods. In the clinical context, "
    "RTX targets the B-cell compartment while TAC inhibits T-cell activation through "
    "calcineurin-NFAT pathway blockade - our cell-type mapping confirmed this complementarity, "
    "with RTX enriched in B cell (89%) and NK cell (33%) markers, and TAC in T cell (11%) markers.",

    "Our network analysis captures RTX's multi-dimensional therapeutic mechanism extending "
    "well beyond B-cell depletion. The recent discovery that RTX induces Treg expansion "
    "through NK cell-mediated TGF{\\beta} production represents a paradigm shift in understanding "
    "anti-CD20 therapy in autoimmune disease [11]. This mechanism explains several clinical "
    "observations: prolonged remission after B-cell repopulation [10], the lack of correlation "
    "between B-cell depletion extent and clinical response [8], and RTX's effects on T-cell "
    "subsets and cytokine profiles [29].",

    "Several limitations should be acknowledged. First, our PPI network may contain "
    "false-positive and false-negative interactions. Second, drug targets were assembled "
    "from multiple sources and may not capture all in vivo targets. Third, our transcriptomic "
    "analysis was based on bulk glomerular tissue. Fourth, the proximity and separation "
    "metrics are correlative and do not establish causality. Nevertheless, our framework "
    "offers advantages over traditional network pharmacology: quantitative assessment, "
    "statistical significance testing, multi-data-type integration, and predictive capability "
    "for drug combination synergy.",

    "In conclusion, this study provides the first systematic network medicine characterization "
    "of rituximab's therapeutic mechanisms in PMN. We demonstrate that RTX's clinical efficacy "
    "arises from its unique ability to target the PMN disease module through multiple coordinated "
    "pathways (z = -5.49). The Complementary Exposure pattern between RTX and TAC (s = 1.12) "
    "provides a quantitative network rationale for their synergistic combination. Our multi-layer "
    "framework offers a generalizable approach for precision medicine in autoimmune kidney diseases."
]

for text in discussion_paragraphs:
    p = doc.add_paragraph(text)
    p.paragraph_format.first_line_indent = Cm(0.75)

# ================================================================
# References
# ================================================================
doc.add_page_break()
doc.add_heading('References', level=1)

references = [
    "[1] Ronco P, Beck L, Debiec H, et al. Membranous nephropathy. Nat Rev Dis Primers. 2021;7:69.",
    "[2] Couser WG. Primary membranous nephropathy. Clin J Am Soc Nephrol. 2017;12:983-97.",
    "[3] Beck LH, Bonegio RG, Lambeau G, et al. M-type phospholipase A2 receptor as target antigen in idiopathic membranous nephropathy. N Engl J Med. 2009;361:11-21.",
    "[4] Francis JM, Beck LH, Salant DJ. Membranous nephropathy: a journey from bench to bedside. Am J Kidney Dis. 2016;68:645-57.",
    "[5] Fervenza FC, Appel GB, Barbour SJ, et al. Rituximab or cyclosporine in the treatment of membranous nephropathy. N Engl J Med. 2019;381:36-46.",
    "[6] Scolari F, Delbarba E, Santoro D, et al. Rituximab or cyclophosphamide in the treatment of membranous nephropathy: the RI-CYCLO randomized trial. J Am Soc Nephrol. 2021;32:972-82.",
    "[7] KDIGO 2021 Clinical Practice Guideline for the Management of Glomerular Diseases. Kidney Int. 2021;100:S1-S276.",
    "[8] Ruggenenti P, Cravedi P, Chianca A, et al. Rituximab in idiopathic membranous nephropathy. J Am Soc Nephrol. 2012;23:1416-25.",
    "[9] Fervenza FC, Abraham RS, Erickson SB, et al. Rituximab therapy in idiopathic membranous nephropathy: a 2-year study. Clin J Am Soc Nephrol. 2010;5:2188-98.",
    "[10] Liu W, Gao C, Liu Z, et al. Rituximab may affect T lymphocyte subsets balance in primary membranous nephropathy. BMC Nephrol. 2024;25:103.",
    "[11] Fribourg M, Canaud G, et al. Rituximab counteracts loss of tolerance in membranous nephropathy patients through NK-mediated Treg induction. medRxiv. 2025.",
    "[12] Weiner GJ. Rituximab: mechanism of action. Semin Hematol. 2010;47:115-23.",
    "[13] Chen Y, Wang J, Liang Q, et al. Clinical study, network pharmacology, and molecular docking of Kunxian capsule in treating idiopathic membranous nephropathy. Front Med. 2025;12:1506972.",
    "[14] Yu H, Li F, Yu H, et al. Exploring the mechanism of Kemofang in treating idiopathic membranous nephropathy. Sci Rep. 2026;16:1985.",
    "[15] Barab\\'asi AL, Gulbahce N, Loscalzo J. Network medicine: a network-based approach to human disease. Nat Rev Genet. 2011;12:56-68.",
    "[16] Menche J, Sharma A, Cho MH, et al. A division of the human proteome reveals disease-disease relationships. Science. 2015;347:1257601.",
    "[17] Guney E, Menche J, Vidal M, Barab\\'asi AL. Network-based in silico drug efficacy screening. Nat Commun. 2016;7:10331.",
    "[18] Cheng F, Desai RJ, Handy DE, et al. Network-based approach to prediction and prevention of drug side effects. Nat Commun. 2019;10:191.",
    "[19] Cai N, Zhu SY, Huang JJ, et al. Rituximab, tacrolimus, cyclophosphamide and cyclosporin in primary membranous nephropathy: network meta-analysis. Int Urol Nephrol. 2025;57:3733-50.",
    "[20] Eddy S, Mariani LH, Kretzler M. Integrated multi-omics approaches to understand glomerular disease. Nephron. 2020;144:166-76.",
    "[21] Ritchie ME, Phipson B, Wu D, et al. limma powers differential expression analyses. Nucleic Acids Res. 2015;43:e47.",
    "[22] Xie J, Liu L, Mladkova N, et al. The genetic architecture of membranous nephropathy. Nat Commun. 2020;11:1600.",
    "[23] Szklarczyk D, Gable AL, Nastou KC, et al. The STRING database in 2021. Nucleic Acids Res. 2021;49:D605-12.",
    "[24] Wishart DS, Feunang YD, Guo AC, et al. DrugBank 5.0. Nucleic Acids Res. 2018;46:D1074-82.",
    "[25] Szklarczyk D, Santos A, von Mering C, et al. STITCH 5. Nucleic Acids Res. 2016;44:D380-4.",
    "[26] Rouillard AD, Gundersen GW, Fernandez NF, et al. The harmonizome. Database. 2016;2016:baw100.",
    "[27] Cheng F, Kov\\'acs IA, Barab\\'asi AL. Network-based prediction of drug combinations. Nat Commun. 2019;10:1197.",
    "[28] Chen EY, Tan CM, Kou Y, et al. Enrichr: interactive enrichment analysis tool. BMC Bioinformatics. 2013;14:128.",
    "[29] Stroopinsky D, Katz T, Rowe JM, et al. Rituximab-induced direct inhibition of T-cell activation. Cancer Immunol Immunother. 2012;61:1233-41.",
]

for ref in references:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(ref)
    run.font.size = Pt(9)

# ================================================================
# Supplementary Figure
# ================================================================
doc.add_page_break()
doc.add_heading('Supplementary Materials', level=1)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_path = os.path.join(FIGURES, "Supplementary_Figure_CellType_Mapping.png")
if os.path.exists(fig_path):
    run = p.add_run()
    run.add_picture(fig_path, width=Inches(4.5))
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Supplementary Figure S1. ')
run.font.bold = True
run = p.add_run('Cell-type mapping of drug targets using scRNA-seq.')
run.font.italic = True

# ================================================================
# Save
# ================================================================
output_path = os.path.join(BASE, "manuscript.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.0f} KB")
