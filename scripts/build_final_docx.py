#!/usr/bin/env python3
"""Generate a Word manuscript directly from manuscript.md.

This script intentionally treats manuscript.md as the single source of truth to
avoid divergence between the Markdown manuscript, DOCX manuscript and reference
export. It supports headings, bold metadata labels, simple bullet lists, figure
legend paragraphs and numbered references.
"""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt, RGBColor

BASE = Path(__file__).resolve().parents[1]
SOURCE = BASE / "manuscript.md"
OUTPUT = BASE / "manuscript.docx"


def configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(6)

    for level in range(1, 4):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 51, 102)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_runs_from_markdown(paragraph, text: str) -> None:
    """Add text with minimal support for **bold** spans."""
    # Strip markdown escapes that are not useful in Word output.
    text = text.replace("\\-", "-")
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)


def add_paragraph(doc: Document, text: str, *, indent: bool = True, bullet: bool = False, center: bool = False) -> None:
    if bullet:
        p = doc.add_paragraph(style="List Bullet")
        add_runs_from_markdown(p, text)
        return
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    add_runs_from_markdown(p, text)


def main() -> None:
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    configure_styles(doc)

    lines = text.splitlines()
    in_references = False
    pending_para: list[str] = []

    def flush() -> None:
        nonlocal pending_para
        if pending_para:
            paragraph_text = " ".join(pending_para).strip()
            add_paragraph(doc, paragraph_text, indent=not in_references)
            pending_para = []

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        if line.strip() == "---":
            flush()
            continue

        h = re.match(r"^(#{1,3})\s+(.*)$", line)
        if h:
            flush()
            level = len(h.group(1))
            heading = h.group(2).strip()
            if level == 1:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(heading)
                r.bold = True
                r.font.size = Pt(16)
                r.font.name = "Times New Roman"
                r.font.color.rgb = RGBColor(0, 51, 102)
            else:
                doc.add_heading(heading, level=level - 1)
            in_references = heading == "References"
            continue

        if line.startswith("- "):
            flush()
            add_paragraph(doc, line[2:].strip(), bullet=True)
            continue

        if re.match(r"^\d+\.\s+", line):
            flush()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line)
            run.font.size = Pt(9)
            continue

        # Metadata fields and legends should not be first-line indented.
        if line.startswith("**") and ":**" in line:
            flush()
            add_paragraph(doc, line, indent=False)
            continue

        pending_para.append(line.strip())

    flush()
    doc.save(OUTPUT)
    print(f"Saved {OUTPUT}")
    print(f"Size: {OUTPUT.stat().st_size / 1024:.0f} KB")


if __name__ == "__main__":
    main()
