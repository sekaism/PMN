#!/usr/bin/env python3
"""Embed figures - clean approach."""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from lxml import etree
import os

BASE = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
FIGURES = os.path.join(BASE, "figures")
SUBMISSION_FIGURES = os.path.join(FIGURES, "submission")
TABLES = os.path.join(BASE, "tables")
DOCX_PATH = os.path.join(BASE, "manuscript.docx")

doc = Document(DOCX_PATH)
body = doc.element.body

def make_figure_para(figname, width, caption_text):
    """Create (fig_para_element, cap_para_element) for a figure."""
    # Figure paragraph
    fig_p = doc.add_paragraph()
    fig_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fig_p.add_run()
    fig_path = os.path.join(FIGURES, figname)
    if os.path.exists(fig_path):
        run.add_picture(fig_path, width=Inches(width))
    
    # Caption paragraph
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap_p.add_run(caption_text)
    r.font.italic = True
    r.font.size = Pt(9)
    
    return fig_p._element, cap_p._element

# Map headings to (figname, width, caption) pairs.
# Figures are condensed into multi-panel figures for the BIB submission draft.
figure_map = {
    "1.3 Study rationale and generalizable contribution": [
        ("submission/Figure1_Workflow_DiseaseModule.png", 6.4,
         "Figure 1. Multi-layer biologic representation workflow and PMN disease-module construction. (A) Workflow integrating glomerular differential expression, PMN GWAS susceptibility genes, STRING PPI expansion, evidence-layered drug-target curation and cell-type marker annotation. (B) Protein-protein interaction network of the PMN disease module. The workflow highlights layered biologic representation and degree-binned robustness checks used to distinguish stable findings from topology-sensitive hypotheses.")
    ],
    "3.3 Rituximab proximity is robust to representation and interactome-sensitivity analyses": [
        ("submission/Figure2_NetworkMetrics.png", 6.4,
         "Figure 2. Network proximity, robustness hierarchy and drug-target separation. (A) Primary size-matched drug-disease proximity z-scores for rituximab, tacrolimus and cyclophosphamide. (B) Pairwise drug-target separation scores. (C) Primary-model putative Complementary Exposure configuration. Degree-binned sensitivity analysis preserved rituximab proximity to the GWAS core and integrated module but did not preserve tacrolimus proximity, so the rituximab-tacrolimus configuration is interpreted as hypothesis-generating.")
    ],
    "3.5 Cell-type marker mapping supports non-redundant immune compartment coverage": [
        ("submission/Figure3_CellType_Mechanism.png", 6.4,
         "Figure 3. Cell-type marker annotation and working mechanism model. (A) Marker-overlap percentages between expanded drug target sets and curated kidney or immune cell markers. (B) Working model of rituximab action through direct B-cell depletion, Fc/complement effector recruitment and putative immune tolerance remodeling. Marker overlap is interpreted as cellular annotation of target distributions rather than as expression-level perturbation.")
    ],
}

table_map = {
    "2.1 Transcriptomic data acquisition and preprocessing": [
        ("Table1_Datasets_and_resources.csv", "Table 1. Datasets and resources. This table lists the transcriptomic, genetic, interactome, drug-target, single-cell and structural resources used in the workflow, together with their roles in each analysis layer.")
    ],
    "2.3 Evidence-layered biologic and comparator drug representation": [
        ("Table2_Target_sets_and_layers.csv", "Table 2. Disease and drug target sets. This table distinguishes seed counts, expanded target counts and final in-network counts after STRING mapping and LCC filtering. For biologic therapy, target layers are interpreted by evidence category rather than as a flat undifferentiated target list.")
    ],
    "3.3 Rituximab proximity is robust to representation and interactome-sensitivity analyses": [
        ("Table3_Network_metrics.csv", "Table 3. Primary network metrics. This table reports primary size-matched drug-disease proximity for DEG-only, GWAS-only and integrated disease sets, together with drug-drug separation results.")
    ],
}

def make_table_elements(csv_name, caption_text):
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_cap.add_run(caption_text)
    r.font.italic = True
    r.font.size = Pt(9)

    import csv
    with open(os.path.join(TABLES, csv_name), newline="") as f:
        rows = list(csv.reader(f))
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = val
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)
                    if i == 0:
                        run.bold = True
    return p_cap._element, table._element

# Find heading paragraphs in the document
heading_elements = []
for para in doc.paragraphs:
    text = para.text.strip()
    for heading in figure_map:
        if heading == text:  # exact match
            heading_elements.append((para._element, figure_map[heading], "figure"))
    for heading in table_map:
        if heading == text:
            heading_elements.append((para._element, table_map[heading], "table"))

# Process from bottom to top to avoid index issues
heading_elements.reverse()

for heading_el, item_list, kind in heading_elements:
    parent = body
    insert_after = heading_el
    
    for item in item_list:
        if kind == "figure":
            figname, width, caption = item
            fig_el, cap_el = make_figure_para(figname, width, caption)
            idx = list(parent).index(insert_after)
            parent.insert(idx + 1, fig_el)
            parent.insert(idx + 2, cap_el)
            insert_after = cap_el
        else:
            csv_name, caption = item
            cap_el, table_el = make_table_elements(csv_name, caption)
            idx = list(parent).index(insert_after)
            parent.insert(idx + 1, cap_el)
            parent.insert(idx + 2, table_el)
            insert_after = table_el

# Supplementary figures at end
doc.add_page_break()
p = doc.add_paragraph()
p.add_run("Supplementary Material").bold = True
fig_el, cap_el = make_figure_para(
    "submission/Supplementary_Figure_S1_Docking.png",
    5.6,
    "Supplementary Figure S1. Optional docking-based structural plausibility screening for tacrolimus and cyclophosphamide against selected topology-derived hub proteins. These scores are not biochemical affinities, they do not validate target engagement, and they were not used to support the primary network medicine conclusions.",
)
fig_el, cap_el = make_figure_para(
    "Supplementary_Figure_S2_External_Validation.png",
    6.2,
    "Supplementary Figure S2. External transcriptomic validation. The figure summarizes discovery-versus-external log2 fold-change concordance, top-ranked DEG overlap across cohorts, and degree-binned drug-disease proximity to the external FDR < 0.05 DEG+GWAS module. External validation provides partial support because global logFC concordance and external-module proximity were positive, whereas strict FDR-level DEG overlap was not enriched.",
)

doc.save(DOCX_PATH)
print(f"Done. Size: {os.path.getsize(DOCX_PATH) / 1024:.0f} KB")
